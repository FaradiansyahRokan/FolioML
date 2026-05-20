import io
import re
import fitz  # PyMuPDF
from typing import List, Dict, Optional

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import csv as csv_module
except ImportError:
    csv_module = None


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".md", ".csv"}


# ─────────────────────────────────────────────────────────────
# Structured Document Model
# ─────────────────────────────────────────────────────────────

class DocumentSection:
    """Represents a structured section of a document."""
    def __init__(self, title: str = "", level: int = 0, content: str = "",
                 section_type: str = "text", metadata: Dict = None):
        self.title = title
        self.level = level  # 0=root, 1=h1, 2=h2, etc.
        self.content = content
        self.section_type = section_type  # text, table, code, list, image_text
        self.metadata = metadata or {}
    
    def to_dict(self) -> Dict:
        return {
            "title": self.title,
            "level": self.level,
            "content": self.content,
            "section_type": self.section_type,
            "metadata": self.metadata,
        }


class StructuredDocument:
    """A fully parsed document with structure preserved."""
    def __init__(self, filename: str):
        self.filename = filename
        self.title = ""
        self.sections: List[DocumentSection] = []
        self.tables: List[Dict] = []
        self.metadata: Dict = {}
        self.page_count: int = 1
        self.pages: List[Dict] = []
    
    @property
    def full_text(self) -> str:
        parts = []
        for sec in self.sections:
            if sec.title:
                prefix = "#" * max(sec.level, 1) + " "
                parts.append(f"{prefix}{sec.title}")
            if sec.content:
                parts.append(sec.content)
        return "\n\n".join(parts)
    
    def to_dict(self) -> Dict:
        return {
            "title": self.title,
            "filename": self.filename,
            "sections": [s.to_dict() for s in self.sections],
            "tables": self.tables,
            "metadata": self.metadata,
            "page_count": self.page_count,
        }


# ─────────────────────────────────────────────────────────────
# Main extraction entry point
# ─────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> Dict:
    """
    Extract text from a file. Returns dict with 'text', 'pages', 'page_count',
    and 'structured' (a StructuredDocument).
    Supports PDF, TXT, DOCX, Markdown, and CSV.
    """
    ext = _get_extension(filename)

    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    elif ext == ".txt":
        return _extract_txt(file_bytes, filename)
    elif ext == ".docx":
        return _extract_docx(file_bytes, filename)
    elif ext == ".md":
        return _extract_markdown(file_bytes, filename)
    elif ext == ".csv":
        return _extract_csv(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _get_extension(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


# ─────────────────────────────────────────────────────────────
# PDF Extraction (enhanced with structure detection)
# ─────────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes, filename: str) -> Dict:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    structured = StructuredDocument(filename)
    pages = []
    all_tables = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        
        # Extract text blocks with position info for structure detection
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        page_text_parts = []
        
        for block in blocks:
            if block["type"] == 0:  # text block
                for line in block.get("lines", []):
                    line_text = ""
                    max_font_size = 0
                    is_bold = False
                    for span in line.get("spans", []):
                        line_text += span["text"]
                        max_font_size = max(max_font_size, span["size"])
                        if "bold" in span.get("font", "").lower() or "Bold" in span.get("font", ""):
                            is_bold = True
                    
                    line_text = line_text.strip()
                    if not line_text:
                        continue
                    
                    # Detect headings by font size
                    if max_font_size >= 16 and len(line_text) < 200:
                        structured.sections.append(DocumentSection(
                            title=line_text, level=1, section_type="heading", metadata={"page": page_num + 1}
                        ))
                    elif max_font_size >= 13 and is_bold and len(line_text) < 200:
                        structured.sections.append(DocumentSection(
                            title=line_text, level=2, section_type="heading", metadata={"page": page_num + 1}
                        ))
                    else:
                        # Append to current section or create new
                        if structured.sections and structured.sections[-1].section_type != "heading":
                            structured.sections[-1].content += "\n" + line_text
                            if "page" not in structured.sections[-1].metadata:
                                structured.sections[-1].metadata["page"] = page_num + 1
                        else:
                            structured.sections.append(DocumentSection(
                                content=line_text, section_type="text", metadata={"page": page_num + 1}
                            ))
                    
                    page_text_parts.append(line_text)
            
            elif block["type"] == 1:  # image block
                # Note: we record that an image exists but can't OCR without Tesseract
                structured.sections.append(DocumentSection(
                    content="[Image detected on page]",
                    section_type="image_text",
                    metadata={"page": page_num + 1}
                ))
        
        # Try to extract tables
        try:
            table_data = page.find_tables()
            if table_data and table_data.tables:
                for table in table_data.tables:
                    extracted = table.extract()
                    if extracted:
                        # Convert table to markdown
                        md_table = _table_to_markdown(extracted)
                        structured.tables.append({
                            "page": page_num + 1,
                            "data": extracted,
                            "markdown": md_table,
                        })
                        structured.sections.append(DocumentSection(
                            content=md_table,
                            section_type="table",
                            metadata={"page": page_num + 1}
                        ))
        except Exception:
            pass  # table extraction is best-effort
        
        page_text = "\n".join(page_text_parts)
        if page_text.strip():
            pages.append({"page": page_num + 1, "text": page_text})

    doc.close()
    
    structured.pages = pages
    structured.page_count = len(pages) or 1
    structured.title = _detect_title(structured)
    structured.metadata = {"format": "pdf", "tables_found": len(all_tables)}
    
    full_text = "\n\n".join(p["text"] for p in pages)
    return {
        "text": full_text,
        "pages": pages,
        "page_count": structured.page_count,
        "structured": structured,
    }


# ─────────────────────────────────────────────────────────────
# TXT Extraction
# ─────────────────────────────────────────────────────────────

def _extract_txt(file_bytes: bytes, filename: str) -> Dict:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    structured = StructuredDocument(filename)
    
    # Try to detect structure from plain text
    lines = text.split("\n")
    current_content = []
    
    for line in lines:
        stripped = line.strip()
        # Detect all-caps headings or lines followed by === or ---
        if stripped and stripped.isupper() and len(stripped) < 100 and len(stripped) > 3:
            if current_content:
                structured.sections.append(DocumentSection(
                    content="\n".join(current_content), section_type="text"
                ))
                current_content = []
            structured.sections.append(DocumentSection(
                title=stripped.title(), level=1, section_type="heading"
            ))
        else:
            current_content.append(line)
    
    if current_content:
        structured.sections.append(DocumentSection(
            content="\n".join(current_content), section_type="text"
        ))
    
    structured.title = _detect_title(structured) or filename
    structured.pages = [{"page": 1, "text": text}]
    structured.page_count = 1
    
    return {"text": text, "pages": [{"page": 1, "text": text}], "page_count": 1, "structured": structured}


# ─────────────────────────────────────────────────────────────
# DOCX Extraction (enhanced with heading detection)
# ─────────────────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes, filename: str) -> Dict:
    if DocxDocument is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")
    
    doc = DocxDocument(io.BytesIO(file_bytes))
    structured = StructuredDocument(filename)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = (para.style.name or "").lower() if para.style else ""
        
        # Detect headings from style
        if "heading 1" in style_name or "title" in style_name:
            structured.sections.append(DocumentSection(title=text, level=1, section_type="heading"))
        elif "heading 2" in style_name:
            structured.sections.append(DocumentSection(title=text, level=2, section_type="heading"))
        elif "heading 3" in style_name:
            structured.sections.append(DocumentSection(title=text, level=3, section_type="heading"))
        elif "code" in style_name:
            structured.sections.append(DocumentSection(content=text, section_type="code"))
        else:
            # Append to last text section if it exists
            if structured.sections and structured.sections[-1].section_type == "text":
                structured.sections[-1].content += "\n\n" + text
            else:
                structured.sections.append(DocumentSection(content=text, section_type="text"))
    
    # Extract tables from DOCX
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            md_table = _table_to_markdown(rows)
            structured.tables.append({"data": rows, "markdown": md_table})
            structured.sections.append(DocumentSection(
                content=md_table, section_type="table"
            ))
    
    structured.title = _detect_title(structured) or filename
    full_text = structured.full_text
    structured.pages = [{"page": 1, "text": full_text}]
    structured.page_count = 1
    
    return {"text": full_text, "pages": [{"page": 1, "text": full_text}], "page_count": 1, "structured": structured}


# ─────────────────────────────────────────────────────────────
# Markdown Extraction
# ─────────────────────────────────────────────────────────────

def _extract_markdown(file_bytes: bytes, filename: str) -> Dict:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    structured = StructuredDocument(filename)
    
    lines = text.split("\n")
    current_content = []
    in_code_block = False
    code_content = []
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                structured.sections.append(DocumentSection(
                    content="\n".join(code_content), section_type="code"
                ))
                code_content = []
                in_code_block = False
            else:
                if current_content:
                    structured.sections.append(DocumentSection(
                        content="\n".join(current_content), section_type="text"
                    ))
                    current_content = []
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line)
            continue
        
        # Detect markdown headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            if current_content:
                structured.sections.append(DocumentSection(
                    content="\n".join(current_content), section_type="text"
                ))
                current_content = []
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            structured.sections.append(DocumentSection(
                title=title, level=level, section_type="heading"
            ))
        else:
            current_content.append(line)
    
    if current_content:
        structured.sections.append(DocumentSection(
            content="\n".join(current_content), section_type="text"
        ))
    if code_content:  # unclosed code block
        structured.sections.append(DocumentSection(
            content="\n".join(code_content), section_type="code"
        ))
    
    structured.title = _detect_title(structured) or filename
    structured.pages = [{"page": 1, "text": text}]
    structured.page_count = 1
    
    return {"text": text, "pages": [{"page": 1, "text": text}], "page_count": 1, "structured": structured}


# ─────────────────────────────────────────────────────────────
# CSV Extraction
# ─────────────────────────────────────────────────────────────

def _extract_csv(file_bytes: bytes, filename: str) -> Dict:
    text = file_bytes.decode("utf-8", errors="replace").strip()
    structured = StructuredDocument(filename)
    
    reader = csv_module.reader(io.StringIO(text))
    rows = list(reader)
    
    if not rows:
        structured.title = filename
        structured.pages = [{"page": 1, "text": ""}]
        structured.page_count = 1
        return {"text": "", "pages": [{"page": 1, "text": ""}], "page_count": 1, "structured": structured}

    header = rows[0]
    data_rows = rows[1:]
    num_rows = len(data_rows)
    
    # Add document title section
    structured.sections.append(DocumentSection(
        title=f"Data from {filename}",
        level=1,
        section_type="heading",
    ))

    # Summary section
    summary = f"This spreadsheet contains {num_rows} data rows with {len(header)} columns: {', '.join(header)}."
    structured.sections.append(DocumentSection(
        content=summary, section_type="text"
    ))

    # Cap total chunks at 50 — calculate batch size dynamically
    MAX_TOTAL_BATCHES = 50
    BATCH_SIZE = max(20, -(-num_rows // MAX_TOTAL_BATCHES))  # ceiling division

    if num_rows <= BATCH_SIZE:
        # Small CSV: one table chunk
        md_table = _table_to_markdown(rows)
        structured.tables.append({"data": rows, "markdown": md_table})
        structured.sections.append(DocumentSection(
            content=md_table, section_type="table"
        ))
    else:
        # Large CSV: split into dynamic batches (max 50 total)
        for batch_start in range(0, num_rows, BATCH_SIZE):
            batch_end = min(batch_start + BATCH_SIZE, num_rows)
            batch_rows = [header] + data_rows[batch_start:batch_end]
            md_table = _table_to_markdown(batch_rows)
            structured.tables.append({"data": batch_rows, "markdown": md_table})
            structured.sections.append(DocumentSection(
                title=f"Rows {batch_start + 1}-{batch_end} of {num_rows}",
                level=2,
                section_type="heading",
            ))
            structured.sections.append(DocumentSection(
                content=md_table, section_type="table"
            ))
    
    structured.title = filename
    full_text = structured.full_text
    structured.pages = [{"page": 1, "text": full_text}]
    structured.page_count = 1
    
    return {"text": full_text, "pages": [{"page": 1, "text": full_text}], "page_count": 1, "structured": structured}


# ─────────────────────────────────────────────────────────────
# Semantic Chunking (MAJOR UPGRADE)
# ─────────────────────────────────────────────────────────────

def chunk_text(text: str, pages: List[Dict] = None, chunk_size: int = 800,
               overlap: int = 200, structured: Optional['StructuredDocument'] = None) -> List[Dict]:
    """
    Split text into chunks. If a StructuredDocument is provided, use semantic
    chunking (heading-aware, table-aware, code-aware). Otherwise fall back
    to the classic sliding-window approach.
    """
    if structured and structured.sections:
        return _semantic_chunk(structured, chunk_size)
    else:
        return _sliding_window_chunk(text, pages, chunk_size, overlap)


def _semantic_chunk(structured: 'StructuredDocument', max_chunk_size: int = 800) -> List[Dict]:
    """
    Semantically-aware chunking that respects document structure.
    - Tables are kept as single chunks (never split)
    - Code blocks are kept intact
    - Text is split at heading boundaries
    - Long text sections are split at paragraph boundaries
    """
    chunks = []
    current_heading = ""
    current_heading_level = 0
    
    for section in structured.sections:
        if section.section_type == "heading":
            current_heading = section.title
            current_heading_level = section.level
            continue
        
        content = section.content.strip()
        if not content:
            continue
        
        # Determine page from metadata or default to 1
        page = section.metadata.get("page", 1)
        
        # Tables and code: keep as single chunk (even if large)
        if section.section_type in ("table", "code"):
            chunk_content = content
            if current_heading:
                chunk_content = f"[Section: {current_heading}]\n\n{content}"
            chunks.append({
                "content": chunk_content,
                "page": page,
                "section": current_heading,
                "type": section.section_type,
            })
            continue
        
        # Text sections: split at paragraph boundaries if too long
        prefix = f"[Section: {current_heading}]\n\n" if current_heading else ""
        
        if len(prefix) + len(content) <= max_chunk_size:
            chunks.append({
                "content": prefix + content,
                "page": page,
                "section": current_heading,
                "type": "text",
            })
        else:
            # Split long text at paragraph boundaries
            paragraphs = content.split("\n\n")
            current_chunk = prefix
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                if len(current_chunk) + len(para) + 2 > max_chunk_size and current_chunk.strip():
                    chunks.append({
                        "content": current_chunk.strip(),
                        "page": page,
                        "section": current_heading,
                        "type": "text",
                    })
                    current_chunk = prefix + para + "\n\n"
                else:
                    current_chunk += para + "\n\n"
            
            if current_chunk.strip() and current_chunk.strip() != prefix.strip():
                chunks.append({
                    "content": current_chunk.strip(),
                    "page": page,
                    "section": current_heading,
                    "type": "text",
                })
    
    # Fallback if structured parsing produced no chunks
    if not chunks and structured.full_text.strip():
        return _sliding_window_chunk(
            structured.full_text,
            structured.pages,
            max_chunk_size,
            200
        )
    
    return chunks


def _sliding_window_chunk(text: str, pages: List[Dict] = None,
                           chunk_size: int = 800, overlap: int = 200) -> List[Dict]:
    """Classic sliding-window chunking with overlap (fallback)."""
    if not text.strip():
        return []

    page_map = _build_page_map(text, pages) if pages else None
    chunks = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size

        if end < text_length:
            # Try to split at natural boundaries
            para_break = text.rfind("\n\n", start, end)
            if para_break > start + chunk_size // 2:
                end = para_break
            else:
                sentence_end = max(
                    text.rfind(". ", start, end),
                    text.rfind("! ", start, end),
                    text.rfind("? ", start, end),
                )
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
                else:
                    word_end = text.rfind(" ", start, end)
                    if word_end > start:
                        end = word_end

        chunk_text_content = text[start:end].strip()
        if chunk_text_content:
            page_num = _get_page_for_position(start, page_map) if page_map else 1
            chunks.append({"content": chunk_text_content, "page": page_num})

        start = end - overlap if end - overlap > start else end

    return chunks


# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────

def _table_to_markdown(rows: List[List]) -> str:
    """Convert a 2D list of strings into a Markdown table."""
    if not rows:
        return ""
    # Clean cells
    clean_rows = []
    for row in rows:
        clean_rows.append([str(cell).strip().replace("|", "\\|") if cell else "" for cell in row])
    
    # Header
    header = "| " + " | ".join(clean_rows[0]) + " |"
    separator = "| " + " | ".join(["---"] * len(clean_rows[0])) + " |"
    body = "\n".join("| " + " | ".join(row) + " |" for row in clean_rows[1:])
    
    return f"{header}\n{separator}\n{body}" if body else f"{header}\n{separator}"


def _detect_title(structured: StructuredDocument) -> str:
    """Try to detect the document title from its structure."""
    for section in structured.sections:
        if section.section_type == "heading" and section.level <= 1 and section.title:
            return section.title
    return ""


def _build_page_map(full_text: str, pages: List[Dict]) -> List[Dict]:
    """Build a mapping of character positions to page numbers."""
    if not pages:
        return []
    mapping = []
    pos = 0
    for page_info in pages:
        page_text = page_info["text"]
        idx = full_text.find(page_text, pos)
        if idx >= 0:
            mapping.append({"start": idx, "end": idx + len(page_text), "page": page_info["page"]})
            pos = idx + len(page_text)
    return mapping


def _get_page_for_position(pos: int, page_map: List[Dict]) -> int:
    """Get the page number for a character position."""
    for entry in page_map:
        if entry["start"] <= pos < entry["end"]:
            return entry["page"]
    return page_map[-1]["page"] if page_map else 1
